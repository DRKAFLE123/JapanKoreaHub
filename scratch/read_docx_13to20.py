import docx
import os

docx_path = "public/Minna_no_Nihongo_I_Lessons_13-20.docx"

if not os.path.exists(docx_path):
    print(f"File not found: {docx_path}")
else:
    doc = docx.Document(docx_path)
    print(f"Doc loaded successfully! Total paragraphs: {len(doc.paragraphs)}, Total tables: {len(doc.tables)}")
    
    # Print first 20 paragraphs to inspect structure
    for i, p in enumerate(doc.paragraphs[:30]):
        if p.text.strip():
            print(f"P{i+1}: {p.text.strip()}")
