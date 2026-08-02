import docx
import json
import sys

sys.stdout.reconfigure(encoding='utf-8')

doc = docx.Document("public/Minna_no_Nihongo_I_Lessons_01-12.docx")

data_all = {}

# 24 tables: Lesson 1 is tables[0, 1], Lesson 2 is tables[2, 3], ..., Lesson 12 is tables[22, 23]
for lesson in range(1, 13):
    v_tbl = doc.tables[(lesson - 1) * 2]
    g_tbl = doc.tables[(lesson - 1) * 2 + 1]
    
    vocab_list = []
    for r in v_tbl.rows[1:]:
        cells = [c.text.strip().replace('\n', ' ') for c in r.cells]
        if len(cells) >= 4 and cells[0]:
            vocab_list.append({
                "kana": cells[0],
                "kanji": cells[1],
                "romaji": cells[2],
                "english": cells[3]
            })
    
    grammar_list = []
    for r in g_tbl.rows[1:]:
        cells = [c.text.strip().replace('\n', ' ') for c in r.cells]
        if len(cells) >= 2 and cells[0]:
            grammar_list.append({
                "title": cells[0],
                "explanation": cells[1],
                "examples": cells[2] if len(cells) > 2 else ""
            })
    
    data_all[lesson] = {
        "vocab": vocab_list,
        "grammar": grammar_list
    }

with open("scratch/docx_lessons_01_12.json", "w", encoding="utf-8") as f:
    json.dump(data_all, f, ensure_ascii=False, indent=2)

print("Extracted DOCX data for Lessons 01-12!")
print("\nSummary:")
for l in range(1, 13):
    v = data_all[l]["vocab"]
    g = data_all[l]["grammar"]
    print(f"  Lesson {l}: {len(v)} vocab items, {len(g)} grammar points")

