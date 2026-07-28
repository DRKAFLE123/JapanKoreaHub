import docx
import json
import sys

sys.stdout.reconfigure(encoding='utf-8')

doc = docx.Document("public/Minna_no_Nihongo_I_Lessons_13-20.docx")

lesson_tables = {
    13: (doc.tables[0], doc.tables[1]),
    14: (doc.tables[2], doc.tables[3]),
    15: (doc.tables[4], doc.tables[5]),
    16: (doc.tables[6], doc.tables[7]),
    17: (doc.tables[8], doc.tables[9]),
    18: (doc.tables[10], doc.tables[11]),
    19: (doc.tables[12], doc.tables[13]),
    20: (doc.tables[14], doc.tables[15]),
}

data_all = {}

for lesson in range(13, 21):
    v_tbl, g_tbl = lesson_tables[lesson]
    
    # Extract Vocab
    vocab_list = []
    for r in v_tbl.rows[1:]:
        cells = [c.text.strip().replace('\n', ' ') for c in r.cells]
        if len(cells) >= 3 and cells[0]:
            vocab_list.append({
                "japanese": cells[0],
                "romaji": cells[1],
                "english": cells[2]
            })
            
    # Extract Grammar
    grammar_list = []
    for r in g_tbl.rows[1:]:
        cells = [c.text.strip().replace('\n', ' ') for c in r.cells]
        if len(cells) >= 2 and cells[0]:
            grammar_list.append({
                "title": cells[0],
                "explanation": cells[1],
                "examples": cells[2] if len(cells) > 2 else ""
            })
            
    data_all[lesson] = {
        "vocab": vocab_list,
        "grammar": grammar_list
    }

with open("scratch/docx_lessons_13_20.json", "w", encoding="utf-8") as f:
    json.dump(data_all, f, ensure_ascii=False, indent=2)

print("Saved scratch/docx_lessons_13_20.json successfully!")
print(f"Summary:")
for l in range(13, 21):
    print(f"  Lesson {l}: {len(data_all[l]['vocab'])} vocab, {len(data_all[l]['grammar'])} grammar points")

