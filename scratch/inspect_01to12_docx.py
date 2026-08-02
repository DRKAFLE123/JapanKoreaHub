import docx
import json
import sys

sys.stdout.reconfigure(encoding='utf-8')

doc = docx.Document("public/Minna_no_Nihongo_I_Lessons_01-12.docx")

print(f"Total tables found in Lessons 01-12.docx: {len(doc.tables)}")

for i, tbl in enumerate(doc.tables):
    headers = [c.text.strip() for c in tbl.rows[0].cells] if tbl.rows else []
    print(f"\nTable {i}: {len(tbl.rows)-1} data rows | Headers: {headers}")
    for r in tbl.rows[1:3]:
        cells = [c.text.strip().replace('\n', ' ')[:60] for c in r.cells]
        print(f"  {cells}")
