import docx
import sys
import os
import re

sys.stdout.reconfigure(encoding='utf-8')

doc = docx.Document("public/Minna_no_Nihongo_I_Lessons_13-20.docx")

# docx has 16 tables:
# Table 1: Lesson 13 Vocab
# Table 2: Lesson 13 Grammar
# Table 3: Lesson 14 Vocab
# Table 4: Lesson 14 Grammar
# Table 5: Lesson 15 Vocab
# Table 6: Lesson 15 Grammar
# Table 7: Lesson 16 Vocab
# Table 8: Lesson 16 Grammar
# Table 9: Lesson 17 Vocab
# Table 10: Lesson 17 Grammar
# Table 11: Lesson 18 Vocab
# Table 12: Lesson 18 Grammar
# Table 13: Lesson 19 Vocab
# Table 14: Lesson 19 Grammar
# Table 15: Lesson 20 Vocab
# Table 16: Lesson 20 Grammar

lesson_tables = {
    13: (doc.tables[0], doc.tables[1]),
    14: (doc.tables[2], doc.tables[3]),
    15: (doc.tables[4], doc.tables[5]),
    16: (doc.tables[6], doc.tables[7]),
    17: (doc.tables[8], doc.tables[9]),
    18: (doc.tables[10], doc.tables[11]),
    19: (doc.tables[12], doc.tables[13]),
    20: (doc.tables[14], doc.tables[15]),
}

for lesson in range(13, 21):
    v_tbl, g_tbl = lesson_tables[lesson]
    print(f"\n================ LESSON {lesson} ================")
    print(f"Vocab items count: {len(v_tbl.rows) - 1}")
    print(f"Grammar points count: {len(g_tbl.rows) - 1}")
    
    print("\nSample Vocab (first 3):")
    for r in v_tbl.rows[1:4]:
        cells = [c.text.strip().replace('\n', ' ') for c in r.cells]
        print("  ", cells)
        
    print("\nGrammar points:")
    for r in g_tbl.rows[1:]:
        cells = [c.text.strip().replace('\n', ' ') for c in r.cells]
        print(f"  • {cells[0]}: {cells[1][:60]}...")
