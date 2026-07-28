import sys
import docx
import json

sys.stdout.reconfigure(encoding='utf-8')

doc = docx.Document("public/Minna_no_Nihongo_I_Lessons_13-20.docx")

print(f"Total tables: {len(doc.tables)}")

for t_idx, table in enumerate(doc.tables):
    print(f"\n--- TABLE {t_idx + 1} ({len(table.rows)} rows, {len(table.columns)} cols) ---")
    for r_idx, row in enumerate(table.rows[:4]):
        cells = [c.text.strip().replace('\n', ' ') for c in row.cells]
        print(f"Row {r_idx + 1}: {cells}")

