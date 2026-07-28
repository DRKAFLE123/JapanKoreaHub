import docx
import json
import sys

sys.stdout.reconfigure(encoding='utf-8')

doc = docx.Document("public/Minna_no_Nihongo_I_Lessons_21-25.docx")

# Tables: 0=L21vocab, 1=L21grammar, 2=L22vocab, 3=L22grammar, 4=L23vocab, 5=L23grammar, 6=L24vocab, 7=L24grammar, 8=L25vocab, 9=L25grammar

lesson_tables = {
    21: (doc.tables[0], doc.tables[1]),
    22: (doc.tables[2], doc.tables[3]),
    23: (doc.tables[4], doc.tables[5]),
    24: (doc.tables[6], doc.tables[7]),
    25: (doc.tables[8], doc.tables[9]),
}

data_all = {}

for lesson in range(21, 26):
    v_tbl, g_tbl = lesson_tables[lesson]
    
    vocab_list = []
    for r in v_tbl.rows[1:]:
        cells = [c.text.strip().replace('\n', ' ') for c in r.cells]
        if len(cells) >= 3 and cells[0]:
            vocab_list.append({
                "japanese": cells[0],
                "romaji": cells[1],
                "english": cells[2]
            })
    
    grammar_list = []
    for r in g_tbl.rows[1:]:
        cells = [c.text.strip().replace('\n', ' ') for c in r.cells]
        if len(cells) >= 2 and cells[0]:
            grammar_list.append({
                "title": cells[0],
                "explanation": cells[1],
                "examples": cells[2] if len(cells) > 2 else ""
            })
    
    data_all[lesson] = {
        "vocab": vocab_list,
        "grammar": grammar_list
    }

with open("scratch/docx_lessons_21_25.json", "w", encoding="utf-8") as f:
    json.dump(data_all, f, ensure_ascii=False, indent=2)

print("Extracted DOCX data for Lessons 21-25!")
print("\nSummary:")
for l in range(21, 26):
    v = data_all[l]["vocab"]
    g = data_all[l]["grammar"]
    print(f"  Lesson {l}: {len(v)} vocab, {len(g)} grammar points")

print("\nSample vocab per lesson:")
for l in range(21, 26):
    print(f"\n  Lesson {l}:")
    for item in data_all[l]["vocab"][:5]:
        print(f"    '{item['japanese']}' ({item['romaji']}) = {item['english']}")
